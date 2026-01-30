"""
Document Loader Module
Loads documents from various formats (TXT, PDF, JSON, DOCX, MD)
Extracts text and preserves metadata
"""

import os
import json
import logging
from pathlib import Path
from typing import List, Dict, Tuple
from dataclasses import dataclass

# PDF handling
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX handling
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a loaded document with content and metadata"""
    content: str
    source: str  # file path
    doc_type: str  # "txt", "pdf", "json", "docx", "md"
    filename: str
    size_bytes: int


class DocumentLoader:
    """
    Loads and extracts text from multiple document formats
    
    Supported formats:
    - .txt: Plain text files
    - .pdf: PDF files (requires PyPDF2)
    - .json: JSON files (assumes "content" or "text" field)
    - .docx: Word documents (requires python-docx)
    - .md: Markdown files (treated as plain text)
    """

    def __init__(self):
        """Initialize the document loader"""
        self.supported_extensions = {
            '.txt', '.pdf', '.json', '.docx', '.md'
        }
        logger.info("DocumentLoader initialized")

    def load_from_directory(self, directory: Path) -> List[Document]:
        """
        Load all supported documents from a directory
        
        Args:
            directory: Path to directory containing documents
            
        Returns:
            List of Document objects
        """
        documents = []
        directory = Path(directory)
        
        if not directory.exists():
            logger.warning(f"Directory not found: {directory}")
            return documents
        
        logger.info(f"Loading documents from: {directory}")
        
        # Get all files in directory
        for file_path in directory.rglob("*"):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.supported_extensions:
                    try:
                        doc = self.load_file(file_path)
                        if doc:
                            documents.append(doc)
                    except Exception as e:
                        logger.error(f"Failed to load {file_path}: {str(e)}")
        
        logger.info(f"Successfully loaded {len(documents)} documents")
        return documents

    def load_file(self, file_path: Path) -> Document | None:
        """
        Load a single document file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Document object or None if loading failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        ext = file_path.suffix.lower()
        logger.debug(f"Loading {ext} file: {file_path.name}")
        
        try:
            if ext == '.txt':
                return self._load_txt(file_path)
            elif ext == '.pdf':
                return self._load_pdf(file_path)
            elif ext == '.json':
                return self._load_json(file_path)
            elif ext == '.docx':
                return self._load_docx(file_path)
            elif ext == '.md':
                return self._load_txt(file_path)  # Treat MD as TXT
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return None
        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            return None

    def _load_txt(self, file_path: Path) -> Document:
        """Load plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return Document(
            content=content,
            source=str(file_path),
            doc_type=file_path.suffix[1:],  # Remove the dot
            filename=file_path.name,
            size_bytes=file_path.stat().st_size
        )

    def _load_pdf(self, file_path: Path) -> Document:
        """Load PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 required for PDF support. Install: pip install PyPDF2")
        
        reader = PdfReader(file_path)
        pages = []
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f"--- Page {page_num + 1} ---\n{text}")
        
        content = "\n\n".join(pages)
        
        return Document(
            content=content,
            source=str(file_path),
            doc_type="pdf",
            filename=file_path.name,
            size_bytes=file_path.stat().st_size
        )

    def _load_json(self, file_path: Path) -> Document:
        """Load JSON file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Try common field names for content
        content = None
        if isinstance(data, dict):
            # Handle FAQ format
            if 'faqs' in data and isinstance(data['faqs'], list):
                faq_texts = []
                for faq in data['faqs']:
                    if 'question' in faq and 'answer' in faq:
                        faq_texts.append(f"Q: {faq['question']}\nA: {faq['answer']}")
                content = "\n\n".join(faq_texts)
            else:
                # Try common field names for content
                for key in ['content', 'text', 'body', 'description', 'answer']:
                    if key in data:
                        content = str(data[key])
                        break
            
            # If no common field found, use entire JSON as string
            if content is None:
                content = json.dumps(data, indent=2)
        else:
            content = json.dumps(data, indent=2)
        
        return Document(
            content=content,
            source=str(file_path),
            doc_type="json",
            filename=file_path.name,
            size_bytes=file_path.stat().st_size
        )

    def _load_docx(self, file_path: Path) -> Document:
        """Load DOCX (Word) file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for DOCX support. Install: pip install python-docx")
        
        doc = DocxDocument(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        content = "\n".join(paragraphs)
        
        return Document(
            content=content,
            source=str(file_path),
            doc_type="docx",
            filename=file_path.name,
            size_bytes=file_path.stat().st_size
        )

    @staticmethod
    def validate_file(file_path: Path, max_size_mb: int = 50) -> Tuple[bool, str]:
        """
        Validate if a file can be loaded
        
        Args:
            file_path: Path to the file
            max_size_mb: Maximum file size in MB
            
        Returns:
            Tuple of (is_valid, message)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return False, "File does not exist"
        
        if not file_path.is_file():
            return False, "Path is not a file"
        
        ext = file_path.suffix.lower()
        if ext not in {'.txt', '.pdf', '.json', '.docx', '.md'}:
            return False, f"Unsupported file type: {ext}"
        
        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb > max_size_mb:
            return False, f"File too large ({size_mb:.1f}MB > {max_size_mb}MB)"
        
        return True, "OK"


# Example usage and testing
if __name__ == "__main__":
    # Setup logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Create a loader
    loader = DocumentLoader()
    
    # Example: Load from a directory
    kb_dir = Path("../../knowledge_base")
    if kb_dir.exists():
        docs = loader.load_from_directory(kb_dir)
        print(f"\nLoaded {len(docs)} documents:")
        for doc in docs:
            print(f"  - {doc.filename} ({doc.doc_type}, {doc.size_bytes} bytes)")
            print(f"    Content preview: {doc.content[:100]}...")
    else:
        print(f"Knowledge base directory not found: {kb_dir}")
